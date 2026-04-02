"""
Serviço de exportação para PDF e DOCX
Gera relatórios profissionais formatados
"""
import io
import os
import re
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any
from xml.sax.saxutils import escape

from app.models import Conversation, Message, MediaType, SentimentType

logger = logging.getLogger(__name__)


def sanitize_for_pdf(text: str) -> str:
    """Remove caracteres fora do BMP que não são suportados pelo ReportLab."""
    if not text:
        return ""
    return re.sub(r'[^\u0000-\uFFFF]', '', text)


def _format_timestamp(dt: Optional[datetime]) -> str:
    if not dt:
        return "—"
    return dt.strftime("%d/%m/%Y às %H:%M:%S")


def _sentiment_label(sentiment: Optional[SentimentType]) -> str:
    labels = {
        SentimentType.POSITIVE: "😊 Positivo",
        SentimentType.NEGATIVE: "😔 Negativo",
        SentimentType.NEUTRAL: "😐 Neutro",
        SentimentType.MIXED: "🤔 Misto",
    }
    return labels.get(sentiment, "—") if sentiment else "—"


def _media_type_label(mtype: MediaType) -> str:
    labels = {
        MediaType.IMAGE: "🖼️ Imagem",
        MediaType.AUDIO: "🎵 Áudio",
        MediaType.VIDEO: "🎬 Vídeo",
        MediaType.DOCUMENT: "📄 Documento",
        MediaType.STICKER: "🎭 Sticker",
        MediaType.CONTACT: "👤 Contato",
        MediaType.LOCATION: "📍 Localização",
        MediaType.DELETED: "🗑️ Mensagem Deletada",
    }
    return labels.get(mtype, "📎 Mídia")


class PDFExporter:
    """Gera PDFs profissionais com ReportLab"""

    def generate(
        self,
        conversation: Conversation,
        messages: List[Message],
        options: Dict[str, bool] = None,
    ) -> bytes:
        """Gera o PDF e retorna como bytes"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm, cm
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                HRFlowable, PageBreak
            )
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
        except ImportError:
            raise RuntimeError("ReportLab não instalado. Execute: pip install reportlab")

        opts = options or {}
        buffer = io.BytesIO()

        def safe_para(text, style):
            """Cria Paragraph com texto sanitizado para PDF."""
            return Paragraph(sanitize_for_pdf(escape(str(text))), style)

        # ─── Configuração do documento ─────────────────────────────────
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
            title=f"Transcrição - {conversation.conversation_name or 'Conversa'}",
            author="WhatsApp Insight Transcriber",
        )

        # ─── Estilos ──────────────────────────────────────────────────
        styles = getSampleStyleSheet()
        BRAND_COLOR = colors.HexColor("#6C63FF")
        DARK_COLOR = colors.HexColor("#1a1a2e")
        ACCENT_COLOR = colors.HexColor("#00d4aa")
        LIGHT_BG = colors.HexColor("#f8f9ff")

        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=24,
            textColor=BRAND_COLOR,
            spaceAfter=6,
            fontName="Helvetica-Bold",
        )
        subtitle_style = ParagraphStyle(
            "Subtitle",
            parent=styles["Normal"],
            fontSize=11,
            textColor=colors.HexColor("#666666"),
            spaceAfter=2,
        )
        section_header_style = ParagraphStyle(
            "SectionHeader",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=BRAND_COLOR,
            spaceBefore=12,
            spaceAfter=6,
            fontName="Helvetica-Bold",
        )
        sender_style = ParagraphStyle(
            "Sender",
            parent=styles["Normal"],
            fontSize=9,
            textColor=ACCENT_COLOR,
            fontName="Helvetica-Bold",
        )
        timestamp_style = ParagraphStyle(
            "Timestamp",
            parent=styles["Normal"],
            fontSize=8,
            textColor=colors.HexColor("#999999"),
            alignment=TA_RIGHT,
        )
        message_style = ParagraphStyle(
            "Message",
            parent=styles["Normal"],
            fontSize=10,
            textColor=DARK_COLOR,
            leading=14,
            spaceAfter=4,
        )
        media_style = ParagraphStyle(
            "Media",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#444444"),
            leftIndent=10,
            leading=13,
        )
        metadata_style = ParagraphStyle(
            "Metadata",
            parent=styles["Normal"],
            fontSize=8,
            textColor=colors.HexColor("#888888"),
            leftIndent=10,
        )

        # ─── Construir conteúdo ───────────────────────────────────────
        story = []

        # Cabeçalho
        story.append(safe_para("WhatsApp Insight Transcriber", title_style))
        story.append(safe_para(
            f"Transcricao Completa da Conversa: <b>{sanitize_for_pdf(escape(conversation.conversation_name or 'Conversa'))}</b>",
            subtitle_style
        ))
        story.append(safe_para(
            f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}",
            subtitle_style
        ))
        story.append(HRFlowable(width="100%", thickness=2, color=BRAND_COLOR))
        story.append(Spacer(1, 10))

        # ─── Metadados da Conversa ────────────────────────────────────
        if opts.get("include_statistics", True):
            story.append(safe_para("Informacoes Gerais", section_header_style))

            participants_str = ", ".join(conversation.participants or []) or "—"
            date_range = f"{_format_timestamp(conversation.date_start)} -> {_format_timestamp(conversation.date_end)}"

            info_data = [
                ["Campo", "Valor"],
                ["Participantes", sanitize_for_pdf(escape(participants_str))],
                ["Periodo", sanitize_for_pdf(escape(date_range))],
                ["Total de Mensagens", str(conversation.total_messages)],
                ["Arquivos de Midia", str(conversation.total_media)],
                ["Sentimento Geral", sanitize_for_pdf(escape(_sentiment_label(conversation.sentiment_overall)))],
            ]

            info_table = Table(info_data, colWidths=[5*cm, 12*cm])
            info_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), BRAND_COLOR),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dddddd")),
                ("BACKGROUND", (0, 1), (-1, -1), LIGHT_BG),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_BG]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(info_table)
            story.append(Spacer(1, 12))

        # ─── Resumo ───────────────────────────────────────────────────
        if opts.get("include_summary", True) and conversation.summary:
            story.append(safe_para("Resumo Executivo", section_header_style))
            story.append(safe_para(conversation.summary, message_style))
            story.append(Spacer(1, 8))

            if conversation.topics:
                topics_str = " - ".join(conversation.topics)
                story.append(safe_para(f"<b>Topicos:</b> {sanitize_for_pdf(escape(topics_str))}", metadata_style))
            story.append(Spacer(1, 12))

        # ─── Contradições ─────────────────────────────────────────────
        if conversation.contradictions and opts.get("include_sentiment_analysis", True):
            story.append(safe_para("Contradicoes Detectadas", section_header_style))
            for c in conversation.contradictions[:10]:
                participant = sanitize_for_pdf(escape(c.get('participant', '?')))
                description = sanitize_for_pdf(escape(c.get('description', '')))
                story.append(safe_para(
                    f"<b>{participant}:</b> {description}",
                    media_style
                ))
            story.append(Spacer(1, 12))

        # ─── Linha do Tempo ───────────────────────────────────────────
        story.append(PageBreak())
        story.append(safe_para("Transcricao Completa", section_header_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#eeeeee")))
        story.append(Spacer(1, 8))

        for msg in messages:
            ts = _format_timestamp(msg.timestamp)

            if msg.media_type == MediaType.TEXT:
                story.append(safe_para(f"{msg.sender}", sender_style))
                story.append(safe_para(ts, timestamp_style))
                story.append(safe_para(msg.original_text or "—", message_style))

            elif msg.media_type == MediaType.DELETED:
                story.append(safe_para(f"{msg.sender}", sender_style))
                story.append(safe_para(ts, timestamp_style))
                story.append(safe_para("<i>Mensagem deletada</i>", metadata_style))

            else:
                story.append(safe_para(f"{msg.sender}", sender_style))
                story.append(safe_para(ts, timestamp_style))

                # Tipo de mídia
                media_label = sanitize_for_pdf(escape(_media_type_label(msg.media_type)))
                media_filename = sanitize_for_pdf(escape(msg.media_filename or ''))
                story.append(safe_para(
                    f"<b>{media_label}</b> -- {media_filename}",
                    media_style
                ))

                # Metadados da mídia
                if msg.media_metadata:
                    meta = msg.media_metadata
                    meta_parts = []
                    if meta.get("file_size_formatted"):
                        meta_parts.append(f"Tamanho: {meta['file_size_formatted']}")
                    if meta.get("duration_formatted"):
                        meta_parts.append(f"Duracao: {meta['duration_formatted']}")
                    if meta.get("resolution"):
                        meta_parts.append(f"Resolucao: {meta['resolution']}")
                    if meta.get("format"):
                        meta_parts.append(f"Formato: {meta['format'].upper()}")
                    if meta_parts:
                        story.append(safe_para(" | ".join(meta_parts), metadata_style))

                # Transcrição/Descrição
                if msg.transcription:
                    story.append(safe_para(f"<b>Transcricao:</b>", media_style))
                    story.append(safe_para(msg.transcription, message_style))
                if msg.description:
                    story.append(safe_para(f"<b>Descricao:</b>", media_style))
                    story.append(safe_para(msg.description, message_style))
                if msg.ocr_text:
                    ocr_text = sanitize_for_pdf(escape(msg.ocr_text))
                    story.append(safe_para(f"<b>Texto (OCR):</b> {ocr_text}", media_style))

            # Sentimento individual
            if opts.get("include_sentiment_analysis", True) and msg.sentiment:
                sentiment_text = sanitize_for_pdf(escape(_sentiment_label(msg.sentiment)))
                story.append(safe_para(
                    f"<i>Sentimento: {sentiment_text}</i>",
                    metadata_style
                ))

            story.append(HRFlowable(
                width="100%", thickness=0.5,
                color=colors.HexColor("#eeeeee"),
                spaceAfter=6
            ))

        # Rodapé
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="100%", thickness=1, color=BRAND_COLOR))
        story.append(safe_para(
            f"<i>Relatorio gerado por WhatsApp Insight Transcriber v1.0 | {datetime.now().strftime('%d/%m/%Y')}</i>",
            metadata_style
        ))

        doc.build(story)
        buffer.seek(0)
        content = buffer.getvalue()
        buffer.close()
        return content


class DOCXExporter:
    """Gera documentos Word profissionais"""

    def generate(
        self,
        conversation: Conversation,
        messages: List[Message],
        options: Dict[str, bool] = None,
    ) -> bytes:
        """Gera o DOCX e retorna como bytes"""
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise RuntimeError("python-docx não instalado. Execute: pip install python-docx")

        opts = options or {}
        doc = Document()

        # ─── Configuração de página ────────────────────────────────────
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # ─── Estilos ──────────────────────────────────────────────────
        def add_heading(text: str, level: int = 1):
            p = doc.add_heading(text, level=level)
            if p.runs:
                run = p.runs[0]
            else:
                run = p.add_run(text)
            run.font.color.rgb = RGBColor(108, 99, 255)
            return p

        def add_paragraph(text: str, bold: bool = False, italic: bool = False, color: str = None, size: int = 10):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            if color:
                r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            return p

        # ─── Cabeçalho ────────────────────────────────────────────────
        title_p = doc.add_paragraph()
        title_run = title_p.add_run("WhatsApp Insight Transcriber")
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(108, 99, 255)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(f"Transcrição: {conversation.conversation_name or 'Conversa'}")
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(100, 100, 100)
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        doc.add_paragraph()  # Espaço

        # ─── Informações Gerais ───────────────────────────────────────
        if opts.get("include_statistics", True):
            add_heading("📊 Informações Gerais", level=2)

            table = doc.add_table(rows=6, cols=2)
            table.style = "Table Grid"

            rows_data = [
                ("Participantes", ", ".join(conversation.participants or []) or "—"),
                ("Data Início", _format_timestamp(conversation.date_start)),
                ("Data Fim", _format_timestamp(conversation.date_end)),
                ("Total Mensagens", str(conversation.total_messages)),
                ("Arquivos de Mídia", str(conversation.total_media)),
                ("Sentimento", _sentiment_label(conversation.sentiment_overall)),
            ]

            for i, (label, value) in enumerate(rows_data):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                # Formatar cabeçalho da célula
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

            doc.add_paragraph()

        # ─── Resumo ───────────────────────────────────────────────────
        if opts.get("include_summary", True) and conversation.summary:
            add_heading("📝 Resumo Executivo", level=2)
            add_paragraph(conversation.summary)
            if conversation.topics:
                add_paragraph("Tópicos: " + " • ".join(conversation.topics), italic=True, size=9)
            doc.add_paragraph()

        # ─── Transcrição ──────────────────────────────────────────────
        doc.add_page_break()
        add_heading("💬 Transcrição Completa", level=1)

        for msg in messages:
            ts = _format_timestamp(msg.timestamp)

            # Remetente e timestamp
            p = doc.add_paragraph()
            run_sender = p.add_run(f"{msg.sender}")
            run_sender.bold = True
            run_sender.font.color.rgb = RGBColor(0, 212, 170)
            run_sender.font.size = Pt(9)
            p.add_run(f"  {ts}")
            p.runs[-1].font.color.rgb = RGBColor(150, 150, 150)
            p.runs[-1].font.size = Pt(8)

            if msg.media_type == MediaType.TEXT:
                add_paragraph(msg.original_text or "—", size=10)

            elif msg.media_type == MediaType.DELETED:
                add_paragraph("🗑️ Mensagem deletada", italic=True, color="999999", size=9)

            else:
                # Tipo de mídia
                p_media = doc.add_paragraph()
                run_type = p_media.add_run(f"{_media_type_label(msg.media_type)}")
                run_type.bold = True
                run_type.font.size = Pt(9)
                if msg.media_filename:
                    p_media.add_run(f" — {msg.media_filename}")

                # Metadados
                if msg.media_metadata:
                    meta = msg.media_metadata
                    meta_parts = []
                    if meta.get("file_size_formatted"):
                        meta_parts.append(f"Tamanho: {meta['file_size_formatted']}")
                    if meta.get("duration_formatted"):
                        meta_parts.append(f"Duração: {meta['duration_formatted']}")
                    if meta.get("resolution"):
                        meta_parts.append(f"Resolução: {meta['resolution']}")
                    if meta_parts:
                        add_paragraph(" | ".join(meta_parts), italic=True, color="888888", size=8)

                # Transcrição/Descrição
                if msg.transcription:
                    add_paragraph("🎤 Transcrição:", bold=True, size=9)
                    add_paragraph(msg.transcription, size=10)
                if msg.description:
                    add_paragraph("👁️ Descrição:", bold=True, size=9)
                    add_paragraph(msg.description, size=10)
                if msg.ocr_text:
                    add_paragraph(f"📄 OCR: {msg.ocr_text}", italic=True, size=9)

            # Linha divisória
            doc.add_paragraph("-" * 60).runs[0].font.size = Pt(6)

        # Rodapé
        doc.add_paragraph()
        footer_p = doc.add_paragraph(
            f"Relatório gerado por WhatsApp Insight Transcriber v1.0 | {datetime.now().strftime('%d/%m/%Y')}"
        )
        for run in footer_p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(150, 150, 150)

        # Salvar para buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        content = buffer.getvalue()
        buffer.close()
        return content


class ExcelExporter:
    """Gera arquivos Excel (.xlsx) com openpyxl"""

    def generate(
        self,
        conversation: Conversation,
        messages: List[Message],
        options: Dict[str, bool] = None,
    ) -> bytes:
        """Gera o XLSX e retorna como bytes"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise RuntimeError("openpyxl não instalado. Execute: pip install openpyxl")

        opts = options or {}
        wb = Workbook()

        # ─── Cores e estilos ─────────────────────────────────────────
        BRAND_COLOR = "6C63FF"
        HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
        HEADER_FILL = PatternFill(start_color=BRAND_COLOR, end_color=BRAND_COLOR, fill_type="solid")
        ALT_FILL_1 = PatternFill(start_color="F8F9FF", end_color="F8F9FF", fill_type="solid")
        ALT_FILL_2 = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        THIN_BORDER = Border(
            left=Side(style="thin", color="DDDDDD"),
            right=Side(style="thin", color="DDDDDD"),
            top=Side(style="thin", color="DDDDDD"),
            bottom=Side(style="thin", color="DDDDDD"),
        )

        def auto_width(ws):
            for col in ws.columns:
                max_length = 0
                col_letter = get_column_letter(col[0].column)
                for cell in col:
                    try:
                        cell_len = len(str(cell.value or ""))
                        if cell_len > max_length:
                            max_length = cell_len
                    except Exception:
                        pass
                ws.column_dimensions[col_letter].width = min(max_length + 4, 60)

        def style_header_row(ws, row_num=1):
            for cell in ws[row_num]:
                cell.font = HEADER_FONT
                cell.fill = HEADER_FILL
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = THIN_BORDER

        def style_data_rows(ws, start_row=2):
            for i, row in enumerate(ws.iter_rows(min_row=start_row, max_row=ws.max_row)):
                fill = ALT_FILL_1 if i % 2 == 0 else ALT_FILL_2
                for cell in row:
                    cell.fill = fill
                    cell.border = THIN_BORDER
                    cell.alignment = Alignment(vertical="top", wrap_text=True)

        # ─── Sheet "Mensagens" ───────────────────────────────────────
        ws_msg = wb.active
        ws_msg.title = "Mensagens"
        ws_msg.append(["Data/Hora", "Remetente", "Mensagem", "Tipo"])
        style_header_row(ws_msg)

        for msg in messages:
            ts = _format_timestamp(msg.timestamp)
            text = msg.original_text or ""
            if msg.transcription:
                text += f"\n[Transcrição] {msg.transcription}"
            if msg.description:
                text += f"\n[Descrição] {msg.description}"
            if msg.ocr_text:
                text += f"\n[OCR] {msg.ocr_text}"
            msg_type = _media_type_label(msg.media_type) if msg.media_type != MediaType.TEXT else "Texto"
            ws_msg.append([ts, msg.sender, text, msg_type])

        style_data_rows(ws_msg)
        auto_width(ws_msg)

        # ─── Sheet "Estatísticas" ────────────────────────────────────
        if opts.get("include_statistics", True):
            ws_stats = wb.create_sheet("Estatísticas")
            ws_stats.append(["Métrica", "Valor"])
            style_header_row(ws_stats)

            ws_stats.append(["Nome da Conversa", conversation.conversation_name or "—"])
            ws_stats.append(["Total de Mensagens", conversation.total_messages])
            ws_stats.append(["Total de Mídias", conversation.total_media])
            ws_stats.append(["Data Início", _format_timestamp(conversation.date_start)])
            ws_stats.append(["Data Fim", _format_timestamp(conversation.date_end)])
            ws_stats.append(["Participantes", ", ".join(conversation.participants or [])])
            ws_stats.append(["Sentimento Geral", _sentiment_label(conversation.sentiment_overall)])

            # Contagem por remetente
            ws_stats.append([])
            ws_stats.append(["Mensagens por Remetente", ""])
            sender_counts: Dict[str, int] = {}
            for msg in messages:
                sender_counts[msg.sender] = sender_counts.get(msg.sender, 0) + 1
            for sender, count in sorted(sender_counts.items(), key=lambda x: x[1], reverse=True):
                ws_stats.append([sender, count])

            # Contagem por tipo
            ws_stats.append([])
            ws_stats.append(["Mensagens por Tipo", ""])
            type_counts: Dict[str, int] = {}
            for msg in messages:
                t = msg.media_type.value
                type_counts[t] = type_counts.get(t, 0) + 1
            for mtype, count in sorted(type_counts.items(), key=lambda x: x[1], reverse=True):
                ws_stats.append([mtype, count])

            style_data_rows(ws_stats)
            auto_width(ws_stats)

        # ─── Sheet "Análise IA" ──────────────────────────────────────
        if opts.get("include_summary", True):
            ws_ai = wb.create_sheet("Análise IA")
            ws_ai.append(["Item", "Conteúdo"])
            style_header_row(ws_ai)

            if conversation.summary:
                ws_ai.append(["Resumo", conversation.summary])
            if conversation.topics:
                ws_ai.append(["Tópicos", ", ".join(conversation.topics)])
            if conversation.keywords:
                ws_ai.append(["Palavras-chave", ", ".join(conversation.keywords)])
            if conversation.contradictions:
                for i, c in enumerate(conversation.contradictions[:20], 1):
                    ws_ai.append([
                        f"Contradição {i}",
                        f"{c.get('participant', '?')}: {c.get('description', '')}"
                    ])
            if conversation.key_moments:
                for i, km in enumerate(conversation.key_moments[:20], 1):
                    ws_ai.append([
                        f"Momento-chave {i}",
                        str(km.get("description", km))
                    ])

            style_data_rows(ws_ai)
            auto_width(ws_ai)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        content = buffer.getvalue()
        buffer.close()
        return content


class CSVExporter:
    """Gera arquivos CSV com encoding UTF-8 BOM"""

    def generate(
        self,
        conversation: Conversation,
        messages: List[Message],
        options: Dict[str, bool] = None,
    ) -> bytes:
        """Gera o CSV e retorna como bytes"""
        import csv

        buffer = io.StringIO()
        writer = csv.writer(buffer, quoting=csv.QUOTE_ALL)

        # Cabeçalho
        writer.writerow(["timestamp", "sender", "message", "type", "is_system"])

        for msg in messages:
            ts = msg.timestamp.isoformat() if msg.timestamp else ""
            text = msg.original_text or ""
            if msg.transcription:
                text += f" [Transcrição: {msg.transcription}]"
            if msg.description:
                text += f" [Descrição: {msg.description}]"
            if msg.ocr_text:
                text += f" [OCR: {msg.ocr_text}]"
            is_system = msg.sender.lower() in ("sistema", "system", "")
            writer.writerow([
                ts,
                msg.sender,
                text,
                msg.media_type.value,
                str(is_system).lower(),
            ])

        csv_text = buffer.getvalue()
        buffer.close()
        # UTF-8 BOM para compatibilidade com Excel
        return b'\xef\xbb\xbf' + csv_text.encode("utf-8")


class HTMLExporter:
    """Gera HTML interativo standalone com estilo WhatsApp"""

    def generate(
        self,
        conversation: Conversation,
        messages: List[Message],
        options: Dict[str, bool] = None,
    ) -> bytes:
        """Gera o HTML e retorna como bytes"""
        opts = options or {}
        conv_name = escape(conversation.conversation_name or "Conversa")
        participants = conversation.participants or []

        # Construir lista de mensagens HTML
        msgs_html = []
        for msg in messages:
            ts = _format_timestamp(msg.timestamp)
            sender = escape(msg.sender)
            text = escape(msg.original_text or "")
            if msg.media_type != MediaType.TEXT:
                media_label = escape(_media_type_label(msg.media_type))
                text = f'<span class="media-badge">{media_label}</span> {text}'
            if msg.transcription:
                text += f'<div class="extra">🎤 {escape(msg.transcription)}</div>'
            if msg.description:
                text += f'<div class="extra">👁️ {escape(msg.description)}</div>'
            if msg.ocr_text:
                text += f'<div class="extra">📄 {escape(msg.ocr_text)}</div>'

            msgs_html.append(
                f'<div class="msg" data-sender="{sender}">'
                f'<div class="msg-header"><span class="sender">{sender}</span>'
                f'<span class="time">{ts}</span></div>'
                f'<div class="msg-body">{text}</div></div>'
            )

        participants_options = "".join(
            f'<option value="{escape(p)}">{escape(p)}</option>' for p in participants
        )

        summary_section = ""
        if opts.get("include_summary", True) and conversation.summary:
            summary_section = f"""
            <div class="section">
                <h2>📝 Resumo</h2>
                <p>{escape(conversation.summary)}</p>
            </div>"""

        stats_section = ""
        if opts.get("include_statistics", True):
            stats_section = f"""
            <div class="section">
                <h2>📊 Estatísticas</h2>
                <table class="stats-table">
                    <tr><td>Total de Mensagens</td><td>{conversation.total_messages}</td></tr>
                    <tr><td>Total de Mídias</td><td>{conversation.total_media}</td></tr>
                    <tr><td>Participantes</td><td>{escape(', '.join(participants))}</td></tr>
                    <tr><td>Período</td><td>{_format_timestamp(conversation.date_start)} — {_format_timestamp(conversation.date_end)}</td></tr>
                    <tr><td>Sentimento</td><td>{escape(_sentiment_label(conversation.sentiment_overall))}</td></tr>
                </table>
            </div>"""

        html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Transcrição - {conv_name}</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Segoe UI',Tahoma,Geneva,Verdana,sans-serif;background:#e5ddd5;color:#1a1a2e}}
.container{{max-width:800px;margin:0 auto;padding:20px}}
.header{{background:linear-gradient(135deg,#6C63FF,#4834d4);color:#fff;padding:30px;border-radius:12px;margin-bottom:20px;text-align:center}}
.header h1{{font-size:1.5em;margin-bottom:5px}}
.header p{{opacity:.8;font-size:.9em}}
.controls{{background:#fff;padding:15px;border-radius:8px;margin-bottom:15px;display:flex;gap:10px;flex-wrap:wrap;align-items:center;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
.controls input,.controls select{{padding:8px 12px;border:1px solid #ddd;border-radius:6px;font-size:14px;flex:1;min-width:150px}}
.controls input:focus,.controls select:focus{{outline:none;border-color:#6C63FF}}
.section{{background:#fff;padding:20px;border-radius:8px;margin-bottom:15px;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
.section h2{{color:#6C63FF;font-size:1.1em;margin-bottom:10px}}
.stats-table{{width:100%;border-collapse:collapse}}
.stats-table td{{padding:8px;border-bottom:1px solid #eee}}
.stats-table td:first-child{{font-weight:bold;width:40%;color:#666}}
.chat{{background:#fff;border-radius:8px;padding:15px;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
.msg{{padding:10px 14px;margin:6px 0;border-radius:8px;background:#dcf8c6;position:relative;max-width:85%}}
.msg:nth-child(even){{background:#fff;border:1px solid #eee}}
.msg.hidden{{display:none}}
.msg-header{{display:flex;justify-content:space-between;align-items:center;margin-bottom:4px}}
.sender{{font-weight:bold;color:#00d4aa;font-size:.85em}}
.time{{color:#999;font-size:.75em}}
.msg-body{{font-size:.95em;line-height:1.5;word-break:break-word}}
.media-badge{{background:#6C63FF;color:#fff;padding:2px 8px;border-radius:4px;font-size:.8em}}
.extra{{margin-top:6px;padding:6px 10px;background:#f0f0f0;border-radius:6px;font-size:.85em;color:#555}}
mark{{background:#ffe066;padding:0 2px;border-radius:2px}}
.footer{{text-align:center;padding:20px;color:#999;font-size:.8em}}
.match-count{{color:#6C63FF;font-weight:bold;font-size:.9em;padding:8px 0}}
</style>
</head>
<body>
<div class="container">
    <div class="header">
        <h1>💬 {conv_name}</h1>
        <p>WhatsApp Insight Transcriber — Exportação Interativa</p>
    </div>
    <div class="controls">
        <input type="text" id="searchBox" placeholder="🔍 Buscar mensagens..." oninput="filterMessages()">
        <select id="senderFilter" onchange="filterMessages()">
            <option value="">Todos os remetentes</option>
            {participants_options}
        </select>
    </div>
    <div class="match-count" id="matchCount"></div>
    {stats_section}
    {summary_section}
    <div class="chat" id="chatContainer">
        {''.join(msgs_html)}
    </div>
    <div class="footer">
        Relatório gerado por WhatsApp Insight Transcriber | {datetime.now().strftime('%d/%m/%Y %H:%M')}
    </div>
</div>
<script>
function filterMessages(){{
    const q=document.getElementById('searchBox').value.toLowerCase();
    const sender=document.getElementById('senderFilter').value;
    const msgs=document.querySelectorAll('.msg');
    let count=0;
    msgs.forEach(m=>{{
        const mSender=m.getAttribute('data-sender');
        const body=m.querySelector('.msg-body');
        const origText=body.getAttribute('data-orig')||body.innerHTML;
        if(!body.getAttribute('data-orig'))body.setAttribute('data-orig',body.innerHTML);
        let show=true;
        if(sender&&mSender!==sender)show=false;
        if(q&&!origText.toLowerCase().includes(q))show=false;
        if(show){{
            m.classList.remove('hidden');
            count++;
            if(q){{
                const re=new RegExp('('+q.replace(/[.*+?^${{}}()|[\\]\\\\]/g,'\\\\$&')+')','gi');
                body.innerHTML=origText.replace(re,'<mark>$1</mark>');
            }}else{{
                body.innerHTML=origText;
            }}
        }}else{{
            m.classList.add('hidden');
        }}
    }});
    document.getElementById('matchCount').textContent=q||sender?count+' mensagem(ns) encontrada(s)':'';
}}
</script>
</body>
</html>"""

        return html.encode("utf-8")


class JSONExporter:
    """Exportação JSON estruturada com metadados completos"""

    def generate(
        self,
        conversation: Conversation,
        messages: List[Message],
        options: Dict[str, bool] = None,
    ) -> bytes:
        """Gera o JSON e retorna como bytes"""
        import json

        opts = options or {}

        data: Dict[str, Any] = {
            "metadata": {
                "exported_at": datetime.now().isoformat(),
                "exporter": "WhatsApp Insight Transcriber v1.0",
                "format_version": "1.0",
            },
            "conversation": {
                "id": conversation.id,
                "name": conversation.conversation_name,
                "participants": conversation.participants or [],
                "total_messages": conversation.total_messages,
                "total_media": conversation.total_media,
                "date_start": conversation.date_start.isoformat() if conversation.date_start else None,
                "date_end": conversation.date_end.isoformat() if conversation.date_end else None,
            },
        }

        if opts.get("include_summary", True):
            data["analysis"] = {
                "summary": conversation.summary,
                "topics": conversation.topics,
                "keywords": conversation.keywords,
                "sentiment_overall": conversation.sentiment_overall.value if conversation.sentiment_overall else None,
                "sentiment_score": conversation.sentiment_score,
                "word_frequency": conversation.word_frequency,
            }

        if opts.get("include_sentiment_analysis", True):
            data["contradictions"] = conversation.contradictions or []
            data["key_moments"] = conversation.key_moments or []

        msg_list = []
        for msg in messages:
            m: Dict[str, Any] = {
                "sequence": msg.sequence_number,
                "timestamp": msg.timestamp.isoformat() if msg.timestamp else None,
                "sender": msg.sender,
                "text": msg.original_text,
                "type": msg.media_type.value,
            }
            if msg.media_filename:
                m["media_filename"] = msg.media_filename
            if msg.media_metadata:
                m["media_metadata"] = msg.media_metadata
            if msg.transcription:
                m["transcription"] = msg.transcription
            if msg.description:
                m["description"] = msg.description
            if msg.ocr_text:
                m["ocr_text"] = msg.ocr_text
            if opts.get("include_sentiment_analysis", True) and msg.sentiment:
                m["sentiment"] = msg.sentiment.value
                m["sentiment_score"] = msg.sentiment_score
            m["is_key_moment"] = msg.is_key_moment
            msg_list.append(m)

        data["messages"] = msg_list

        return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
